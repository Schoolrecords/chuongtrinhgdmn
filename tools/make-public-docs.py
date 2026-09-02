# -*- coding: utf-8 -*-
"""Tạo các tệp Word cho bản DÙNG CHUNG (chuongtrinhgiaoduc.quantrisotruonghoc.com).

1. Chép mọi tệp KHDH trong website/assets/docs/ sang web-chung/assets/docs/, thay tên trường,
   tên người kí, địa danh và ngày tháng bằng dòng chấm để trường nào cũng dùng được.
2. Dựng "Mẫu Phụ lục 2 – KHDH môn học (trống)" từ một tệp thật: giữ nguyên khuôn trang, đầu trang,
   bảng 6 cột, bảng quy ước mã tích hợp; xoá hết nội dung bài học.

Chạy trong thư mục website/:   python tools/make-public-docs.py
Cần: pip install python-docx
"""
import copy
import os
import re
import shutil
import sys

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

HERE = os.path.dirname(os.path.abspath(__file__))
WEBSITE = os.path.dirname(HERE)
OUT = os.path.join(os.path.dirname(WEBSITE), 'web-chung')

DOTS = '……………………………'

# Thay thế theo thứ tự: cụ thể trước, chung sau
REPLACEMENTS = [
    # Bản viết hoa (đầu trang) giữ dạng hoa; các dạng còn lại giữ dạng chữ thường trong câu văn
    (re.compile(r'UBND\s+XÃ\s+QUẢNG\s+CHÂU'), 'UBND ……………………'),
    (re.compile(r'TRƯỜNG\s+TIỂU\s+HỌC\s+QUẢNG\s+CHÂU\s*1'), 'TRƯỜNG TIỂU HỌC ……………………'),
    (re.compile(r'UBND\s+xã\s+Quảng\s+Châu', re.I), 'UBND xã ……………………'),
    (re.compile(r'Trường\s+Tiểu\s+học\s+Quảng\s+Châu\s*1', re.I), 'Trường Tiểu học ……………………'),
    (re.compile(r'Quảng\s+Châu,\s*ngày\s*\d*\s*tháng\s*\d*\s*năm\s*\d*'),
     '……………, ngày … tháng … năm 20…'),
    (re.compile(r'Quảng\s+Châu'), '……………'),
    (re.compile(r'Trần\s+Thị\s+Liên'), ''),
]


def replace_text(s):
    out = s
    for pat, rep in REPLACEMENTS:
        out = pat.sub(rep, out)
    return out


def rewrite_paragraph(p, fn=replace_text):
    """Đổi nội dung một đoạn nhưng giữ định dạng của run đầu tiên."""
    old = p.text
    new = fn(old)
    if new == old:
        return False
    runs = p.runs
    if not runs:
        return False
    runs[0].text = new
    for r in runs[1:]:
        r.text = ''
    return True


def walk_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for sec in doc.sections:
        for part in (sec.header, sec.footer, sec.first_page_header, sec.first_page_footer,
                     sec.even_page_header, sec.even_page_footer):
            for p in part.paragraphs:
                yield p


def anonymize(src, dst):
    doc = Document(src)
    n = 0
    for p in walk_paragraphs(doc):
        if rewrite_paragraph(p):
            n += 1
    os.makedirs(os.path.dirname(dst), exist_ok=True)
    doc.save(dst)
    return n


# ---------------------------------------------------------------- mẫu trống
def body_items(doc):
    """Danh sách (loại, đối tượng) theo đúng thứ tự trong thân tài liệu."""
    items = []
    for child in doc.element.body.iterchildren():
        tag = child.tag.split('}')[1]
        if tag == 'p':
            items.append(('p', Paragraph(child, doc)))
        elif tag == 'tbl':
            items.append(('tbl', Table(child, doc)))
    return items


def set_row_text(row, text):
    for cell in row.cells:
        for i, p in enumerate(cell.paragraphs):
            if i == 0:
                if p.runs:
                    p.runs[0].text = text
                    for r in p.runs[1:]:
                        r.text = ''
            else:
                p._element.getparent().remove(p._element)
        break  # dòng gộp ô: chỉ cần ô đầu


BLANK_HEAD = [
    (re.compile(r'KHỐI\s+LỚP\s*\d'), 'KHỐI LỚP …'),
    (re.compile(r'Năm\s+học\s*20\d\d\s*-\s*20\d\d'), 'Năm học 20… - 20…'),
    (re.compile(r'năm\s+học\s*20\d\d\s*-\s*20\d\d'), 'năm học 20… - 20…'),
    (re.compile(r'lớp\s*\d(?![\d])'), 'lớp …'),
    (re.compile(r'(1\.\s*Môn học, hoạt động giáo dục:).*'), r'\1 ……………………………………'),
]


def blank_head(s):
    out = replace_text(s)
    for pat, rep in BLANK_HEAD:
        out = pat.sub(rep, out)
    return out


def make_blank_template(src, dst, blank_rows=8):
    doc = Document(src)
    items = body_items(doc)

    # Chỉ làm trống lớp/năm học ở phần đầu (trước bảng kế hoạch); phần ghi chú phía sau giữ nguyên
    seen_table = False
    for kind, obj in items:
        if kind == 'tbl':
            seen_table = True
        elif kind == 'p':
            rewrite_paragraph(obj, replace_text if seen_table else blank_head)

    tables = [o for k, o in items if k == 'tbl']
    if not tables:
        raise SystemExit('Không tìm thấy bảng KHDH trong ' + src)
    plan = tables[0]

    rows = plan.rows
    head = 2  # 2 dòng tiêu đề gộp ô
    sem_tr = copy.deepcopy(rows[head]._tr)   # dòng "HỌC KÌ I" gộp cả hàng
    data_tr = copy.deepcopy(rows[head + 1]._tr)

    tbl = plan._tbl
    for tr in list(tbl.tr_lst)[head:]:
        tbl.remove(tr)

    def add(tr, text=None):
        new = copy.deepcopy(tr)
        tbl.append(new)
        row = plan.rows[-1]
        if text is None:
            for cell in row.cells:
                for i, p in enumerate(cell.paragraphs):
                    if i == 0:
                        if p.runs:
                            p.runs[0].text = ''
                            for r in p.runs[1:]:
                                r.text = ''
                    else:
                        p._element.getparent().remove(p._element)
        else:
            set_row_text(row, text)

    for label in ('HỌC KÌ I', 'HỌC KÌ II'):
        add(sem_tr, label)
        for _ in range(blank_rows):
            add(data_tr)

    # Bảng chữ kí: bỏ tên người kí, để dòng chấm
    if len(tables) >= 3:
        for row in tables[-1].rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    rewrite_paragraph(p, blank_head)

    os.makedirs(os.path.dirname(dst), exist_ok=True)
    doc.save(dst)


def main():
    src_docs = os.path.join(WEBSITE, 'assets', 'docs')
    dst_docs = os.path.join(OUT, 'assets', 'docs')
    if os.path.isdir(dst_docs):
        shutil.rmtree(dst_docs)
    count = 0
    for root, _dirs, files in os.walk(src_docs):
        for fn in files:
            if not fn.lower().endswith('.docx'):
                continue
            rel = os.path.relpath(os.path.join(root, fn), src_docs)
            anonymize(os.path.join(root, fn), os.path.join(dst_docs, rel))
            count += 1
    print(f'Đã ẩn danh {count} tệp Word -> {dst_docs}')

    # Khu "Mẫu biểu trống" đã bỏ khỏi trang (2/9/2026). Muốn bật lại: khai báo mục "templates"
    # trong tools/public/school.json rồi bỏ chú thích khối dưới đây.
    tpl_dir = os.path.join(OUT, 'assets', 'templates')
    if os.path.isdir(tpl_dir):
        shutil.rmtree(tpl_dir)
    # base = os.path.join(src_docs, 'lop1', 'Toan - Lop 1.docx')
    # make_blank_template(base, os.path.join(tpl_dir, 'Mau Phu luc 2 - KHDH mon hoc (trong).docx'))


if __name__ == '__main__':
    sys.exit(main())
