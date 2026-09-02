# -*- coding: utf-8 -*-
"""
Nhập Kế hoạch dạy học (KHDH) từ các tệp Word (Phụ lục 2 CV 2345) sang JSON.

Cách chạy (từ thư mục website/):
    python tools/docx-to-json.py              # đọc ../Lớp 1 … ../Lớp 5, ghi vào data/curriculum/
    python tools/docx-to-json.py --src "đường/dẫn/thư mục chứa Lớp 1..5"

Sau khi chạy, chạy tiếp:  node tools/build-data.mjs  để đóng gói JSON thành js/data.bundle.js.

Yêu cầu: pip install python-docx
"""
import argparse
import datetime
import glob
import json
import os
import re
import shutil
import sys

try:
    import docx  # python-docx
except ImportError:
    sys.exit("Thiếu thư viện python-docx. Cài bằng: pip install python-docx")

HERE = os.path.dirname(os.path.abspath(__file__))
WEBSITE = os.path.dirname(HERE)
DEFAULT_SRC = os.path.dirname(WEBSITE)          # thư mục "Chương trình GD Môn học"
OUT_DIR = os.path.join(WEBSITE, "data", "curriculum")
DOCS_DIR = os.path.join(WEBSITE, "assets", "docs")   # tệp Word đính kèm để tải về
SCHOOL_YEAR = "2026-2027"
# Ngày nhập dữ liệu ghi vào JSON. Giữ nguyên ngày nhà trường bàn giao tệp Word,
# để chạy lại bộ nhập liệu (sửa lỗi phân tích) không làm đổi ngày. Đổi bằng --date.
IMPORT_DATE = "2026-09-02"

# Tên môn trong tên tệp -> mã môn dùng trong website (khớp data/subjects.json)
SUBJECT_MAP = [
    ("Tiếng Việt", "tieng-viet"),
    ("Toán", "toan"),
    ("Đạo đức", "dao-duc"),
    ("Tự nhiên và Xã hội", "tu-nhien-va-xa-hoi"),
    ("Khoa học", "khoa-hoc"),
    ("Lịch sử & Địa lí", "lich-su-va-dia-li"),
    ("Lịch sử và Địa lí", "lich-su-va-dia-li"),
    ("Tiếng Anh", "ngoai-ngu-1"),
    ("Ngoại ngữ 1", "ngoai-ngu-1"),
    ("Tin học", "tin-hoc"),
    ("Công nghệ", "cong-nghe"),
    ("Âm nhạc", "am-nhac"),
    ("Mĩ thuật", "mi-thuat"),
    ("GDTC", "giao-duc-the-chat"),
    ("Giáo dục thể chất", "giao-duc-the-chat"),
    ("HĐTN", "hoat-dong-trai-nghiem"),
    ("Hoạt động trải nghiệm", "hoat-dong-trai-nghiem"),
]

# Mã nội dung tích hợp, lồng ghép xuất hiện trong cột "Nội dung điều chỉnh, bổ sung"
INTEGRATION_CODES = [
    "NLS-KT", "NLS-GT", "NLS-ST", "NLS-AT", "NLS-GQVĐ", "NLS-NB", "NLS-ƯD", "NLS",
    "AI-NB", "AI",
    "GDQPAN", "QPAN",
    "KNS", "STEM",
    "GDĐP", "ATGT", "GDMT", "QTE", "ĐĐLS", "GDHN", "GDVL",
]

# Cụm từ tự do (không có mã) trong cột điều chỉnh -> mã tích hợp. So khớp trên chuỗi đã bỏ dấu, chữ thường.
PHRASE_CODES = [
    ("quoc phong", "QPAN"), ("gdqp", "QPAN"), ("qpan", "QPAN"),
    ("dia phuong", "GDĐP"), ("gddp", "GDĐP"),
    ("viet lao", "GDVL"), ("viet - lao", "GDVL"), ("viet – lao", "GDVL"),
    ("stem", "STEM"),
    ("an toan giao thong", "ATGT"), ("atgt", "ATGT"),
    ("moi truong", "GDMT"), ("quyen tre em", "QTE"), ("hoa nhap", "GDHN"),
    ("dao duc, loi song", "ĐĐLS"), ("ki nang song", "KNS"), ("nang luc so", "NLS"),
    ("tri tue nhan tao", "AI-NB"),
]


def _strip(text):
    import unicodedata
    text = unicodedata.normalize("NFD", text).replace("đ", "d").replace("Đ", "D")
    return "".join(ch for ch in text if unicodedata.category(ch) != "Mn").lower()


def phrases_to_integrations(text):
    """Tách phần văn bản tự do thành (điều chỉnh còn lại, [tích hợp nhận diện theo cụm từ])."""
    keep, items = [], []
    for seg in re.split(r"\s*;\s*|\n|\s+[-–—]\s+", text or ""):
        seg = seg.strip(" .,")
        if not seg:
            continue
        n = _strip(seg)
        # Một đoạn có thể ghi nhiều nội dung: "GDQP, GD Việt Lào" -> lấy đủ mã, không dừng ở mã đầu tiên
        codes = []
        for k, c in PHRASE_CODES:
            if k in n and c not in codes:
                codes.append(c)
        # Chỉ coi là tích hợp khi đoạn ngắn, mang tính nhãn (không phải câu điều chỉnh dài)
        if codes and len(seg) <= 90 and not n.startswith("dieu chinh"):
            for c in codes:
                items.append({"code": c, "level": "", "text": seg})
        else:
            keep.append(seg)
    return "; ".join(keep), items

# "ĐÍNH CHÍNH" là điều chỉnh SGK, không phải tích hợp
ADJUST_CODES = ["ĐÍNH CHÍNH", "ĐÍNH CHÍNH SGK", "ĐIỀU CHỈNH"]

CODE_RE = re.compile(
    r"(?:^|(?<=[\.\;\)\s]))(?:Lồng ghép\s+)?(" +
    "|".join(re.escape(c) for c in sorted(INTEGRATION_CODES + ADJUST_CODES, key=len, reverse=True)) +
    r")\s*(\([^)]*\))?\s*([:,])",
    re.UNICODE,
)
PERIOD_RE = re.compile(r"Tiết\s*(\d+)\s*(?:[–\-+]\s*(\d+))?\s*(?:\((\d+)\s*tiết\))?", re.UNICODE)


def subject_id_from_filename(name):
    for key, sid in SUBJECT_MAP:
        if key.lower() in name.lower():
            return sid
    return None


def split_integrations(text):
    """Tách cột 'Nội dung điều chỉnh, bổ sung' thành (điều chỉnh tự do, [tích hợp])."""
    text = (text or "").strip()
    if not text:
        return "", []
    matches = list(CODE_RE.finditer(text))
    if not matches:
        return phrases_to_integrations(text)
    free = text[: matches[0].start()].strip(" ;.\n")
    items = []
    adjust_parts = []
    if free:
        rest, found = phrases_to_integrations(free)
        items.extend(found)
        if rest:
            adjust_parts.append(rest)
    pending = []  # các mã viết liền nhau "NLS-KT, KNS: nội dung" dùng chung nội dung phía sau
    for i, m in enumerate(matches):
        code = m.group(1)
        level = (m.group(2) or "").strip("() ")
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[m.end():end].strip(" ;\n")
        norm = {"GDQPAN": "QPAN", "AI": "AI-NB"}.get(code, code)
        if code in ADJUST_CODES:
            adjust_parts.append(body)
            continue
        if m.group(3) == "," and not body:
            pending.append((norm, level))
            continue
        for pc, pl in pending:
            items.append({"code": pc, "level": pl or level, "text": body})
        pending = []
        items.append({"code": norm, "level": level, "text": body})
    for pc, pl in pending:
        items.append({"code": pc, "level": pl, "text": ""})
    return "; ".join(p for p in adjust_parts if p), items


def parse_period(label):
    """Trả về (tiết đầu, tiết cuối, số tiết của dòng, tổng số tiết của bài nếu ghi trong ngoặc).

    Trong tệp KHDH, một bài nhiều tiết có thể được tách thành nhiều dòng: dòng đầu ghi
    "Tiết 2 (3 tiết)" (3 tiết là tổng của cả bài), các dòng sau chỉ ghi "Tiết 3", "Tiết 4".
    Số tiết của từng dòng luôn tính theo khoảng tiết, không lấy số trong ngoặc.
    """
    m = PERIOD_RE.search(label or "")
    if not m:
        n = re.search(r"(\d+)\s*tiết", label or "")
        n = int(n.group(1)) if n else None
        return None, None, n, n
    a = int(m.group(1))
    b = int(m.group(2)) if m.group(2) else a
    lesson_total = int(m.group(3)) if m.group(3) else None
    return a, b, b - a + 1, lesson_total


ALIGN = {0: "left", 1: "center", 2: "right", 3: "justify"}


def extract_document(d):
    """Lấy phần văn bản ngoài bảng KHDH (đầu trang, mục I–IV, quy ước mã, chữ kí) để dựng trang A4."""
    from docx.text.paragraph import Paragraph
    blocks = []      # các đoạn trước bảng KHDH
    after = []       # các đoạn sau bảng KHDH (trừ bảng)
    legend, signature = [], []
    seen_main = False
    ti = 0
    for child in d.element.body.iterchildren():
        tag = child.tag.split("}")[1]
        if tag == "tbl":
            t = d.tables[ti]; ti += 1
            if ti == 1:
                seen_main = True
                continue
            rows = [[c.text.strip() for c in r.cells] for r in t.rows]
            if len(t.columns) == 3 and rows and rows[0][0].strip().lower() == "mã":
                legend = rows[1:]
            elif len(t.columns) == 2 and len(rows) == 1:
                signature = [[ln.strip() for ln in c.splitlines()] for c in [t.rows[0].cells[0].text, t.rows[0].cells[1].text]]
            else:
                after.append({"table": rows})
            continue
        if tag != "p":
            continue
        para = Paragraph(child, d)
        text = para.text.strip()
        if not text:
            continue
        runs = para.runs
        blk = {
            "text": text,
            "align": ALIGN.get(para.alignment, "left"),
            "bold": any(r.bold for r in runs),
            "italic": any(r.italic for r in runs) and not any(r.bold for r in runs),
        }
        (after if seen_main else blocks).append(blk)
    widths = []
    try:
        widths = [round(c.width.mm) if c.width else None for c in d.tables[0].columns]
    except Exception:  # noqa
        pass
    sec = d.sections[0]
    return {
        "page": {"orientation": "landscape" if sec.page_width > sec.page_height else "portrait",
                 "widthMm": round(sec.page_width.mm), "heightMm": round(sec.page_height.mm),
                 "marginMm": round(sec.left_margin.mm)},
        "columnWidthsMm": widths,
        "before": blocks,
        "legend": legend,
        "after": after,
        "signature": signature,
    }


def parse_docx(path, grade, subject_id):
    d = docx.Document(path)
    if not d.tables:
        raise ValueError("Không có bảng trong tệp")
    table = d.tables[0]
    lessons = []
    semester = 1
    cur_week = None
    cur_theme = ""
    idx = 0
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 6:
            continue
        joined = " ".join(set(cells)).upper()
        if ri < 2 and "TUẦN" in joined:
            continue  # dòng tiêu đề
        # Dòng "HỌC KÌ I/II" gộp toàn bộ cột
        if len(set(cells)) == 1 and "HỌC KÌ" in joined:
            semester = 2 if ("II" in joined.replace("HỌC KÌ", "").strip()) else 1
            continue
        week_txt, theme, title, period_label, adjust_txt, note = cells[:6]
        if week_txt:
            wk = re.search(r"\d+", week_txt)
            if wk:
                cur_week = int(wk.group(0))
        if theme:
            # Ô chủ đề bị gõ tách làm hai dòng trong Word ("CHỦ ĐỀ 4. CÁC NƯỚC LÁNG" / "GIỀNG"):
            # mảnh chỉ có một từ thì nối vào chủ đề "CHỦ ĐỀ/CHỦ ĐIỂM ..." ngay trước đó.
            if (len(theme.split()) == 1 and cur_theme
                    and re.match(r"\s*CHỦ\s+(ĐỀ|ĐIỂM)", cur_theme, re.I)
                    and not cur_theme.endswith(theme)):
                old_theme, cur_theme = cur_theme, f"{cur_theme} {theme}"
                for prev in reversed(lessons):   # vá lại các dòng đã ghi với tên chủ đề bị cắt
                    if prev["theme"] != old_theme:
                        break
                    prev["theme"] = cur_theme
            else:
                cur_theme = theme
        if not title and not period_label:
            continue
        adjustments, integrations = split_integrations(adjust_txt)
        p_from, p_to, p_n, p_lesson = parse_period(period_label)
        idx += 1
        lessons.append({
            "id": f"l{grade}-{subject_id}-{idx:03d}",
            "semester": semester if cur_week is None else (semester if semester == 2 or cur_week <= 18 else 2),
            "week": cur_week,
            "theme": cur_theme,
            "title": title,
            "content": "",
            "periodLabel": period_label,
            "periodFrom": p_from,
            "periodTo": p_to,
            "periods": p_n,
            "lessonPeriods": p_lesson,
            "adjustments": adjustments,
            "integrations": integrations,
            "note": note,
        })
    total = sum(l["periods"] or 0 for l in lessons)
    weeks = sorted({l["week"] for l in lessons if l["week"]})
    return {
        "schoolYear": SCHOOL_YEAR,
        "grade": grade,
        "subjectId": subject_id,
        "status": "draft",
        "statusLabel": "Bản nháp từ tệp KHDH, chờ tổ chuyên môn xác nhận",
        "source": {
            "file": os.path.relpath(path, DEFAULT_SRC).replace("\\", "/"),
            "importedAt": IMPORT_DATE,
        },
        "summary": {
            "totalPeriods": total,
            "weeks": len(weeks),
            "semester1Weeks": len([w for w in weeks if w <= 18]),
            "semester2Weeks": len([w for w in weeks if w > 18]),
            "lessons": len(lessons),
            "integrations": len([1 for l in lessons for _ in l["integrations"]]),
        },
        "lessons": lessons,
        "document": extract_document(d),
    }


def ascii_name(text):
    """'Tiếng Việt' -> 'Tieng Viet' (bỏ dấu, giữ hoa/thường) để đặt tên tệp tải về."""
    import unicodedata
    t = unicodedata.normalize("NFD", text).replace("đ", "d").replace("Đ", "D")
    t = "".join(ch for ch in t if unicodedata.category(ch) != "Mn")
    return re.sub(r"[^A-Za-z0-9 \-]+", "", t).strip()


def subject_display_name(sid):
    """Tên môn để đặt tên tệp, lấy từ data/subjects.json (Ngoại ngữ 1 -> Tiếng Anh)."""
    try:
        with open(os.path.join(WEBSITE, "data", "subjects.json"), encoding="utf-8") as fh:
            for sub in json.load(fh):
                if sub["id"] == sid:
                    return sub.get("subtitle") or sub["name"]
    except Exception:  # noqa
        pass
    return sid


def attach(src, grade, name):
    """Chép tệp Word vào assets/docs/lopN/<Tên môn> - Lop N.docx và trả về mô tả đính kèm."""
    d = os.path.join(DOCS_DIR, f"lop{grade}")
    os.makedirs(d, exist_ok=True)
    fname = f"{name} - Lop {grade}.docx"
    dst = os.path.join(d, fname)
    shutil.copyfile(src, dst)
    st = os.stat(src)
    return {
        "file": f"assets/docs/lop{grade}/{fname}",
        "name": fname,
        "original": os.path.basename(src),
        "size": st.st_size,
        "updated": datetime.date.fromtimestamp(st.st_mtime).isoformat(),
    }


def main():
    global IMPORT_DATE
    ap = argparse.ArgumentParser()
    ap.add_argument("--src", default=DEFAULT_SRC)
    ap.add_argument("--date", default=IMPORT_DATE, help="Ngày nhập ghi vào JSON (mặc định %(default)s)")
    args = ap.parse_args()
    IMPORT_DATE = args.date
    os.makedirs(OUT_DIR, exist_ok=True)
    if os.path.isdir(DOCS_DIR):   # xoá các tệp cũ để không còn tệp thừa (giữ thư mục vì Google Drive khoá thư mục)
        for root, _dirs, fnames in os.walk(DOCS_DIR):
            for fn in fnames:
                if fn.lower().endswith(".docx"):
                    os.remove(os.path.join(root, fn))
    index = []
    files = sorted(glob.glob(os.path.join(args.src, "Lớp *", "KHDH *lớp * (*).docx")))
    grade_files = sorted(glob.glob(os.path.join(args.src, "Lớp *", "KHDH cả khối * (*).docx")))
    files = [f for f in files if "cả khối" not in os.path.basename(f)]
    if not files:
        sys.exit(f"Không tìm thấy tệp KHDH trong {args.src}")
    # Mỗi (lớp, môn) chỉ lấy một tệp: ưu tiên bản "(đã rà soát)" nếu có
    chosen = {}
    for f in files:
        base = os.path.basename(f)
        g = re.search(r"lớp\s*(\d)", base, re.I)
        sid = subject_id_from_filename(base.split("lớp")[0])
        if not g or not sid:
            print("BỎ QUA (không nhận diện được):", base)
            continue
        key = (int(g.group(1)), sid)
        reviewed = "rà soát" in base.lower()
        if key not in chosen or (reviewed and not chosen[key][1]):
            chosen[key] = (f, reviewed)
    for (grade, sid), (f, reviewed) in sorted(chosen.items()):
        base = os.path.basename(f)
        try:
            data = parse_docx(f, grade, sid)
        except Exception as e:  # noqa
            print("LỖI", base, e)
            continue
        if reviewed:
            data["status"] = "reviewed"
            data["statusLabel"] = "Đã rà soát (bản tổ chuyên môn rà soát ngày 02/9/2026), chờ Hiệu trưởng phê duyệt"
        data["attachment"] = attach(f, grade, ascii_name(subject_display_name(sid)))
        out_dir = os.path.join(OUT_DIR, f"lop{grade}")
        os.makedirs(out_dir, exist_ok=True)
        out = os.path.join(out_dir, f"{sid}.json")
        with open(out, "w", encoding="utf-8") as fh:
            json.dump(data, fh, ensure_ascii=False, indent=1)
        index.append({
            "grade": grade, "subjectId": sid, "file": f"curriculum/lop{grade}/{sid}.json",
            "status": data["status"], "summary": data["summary"], "source": data["source"],
            "attachment": data["attachment"],
        })
        s = data["summary"]
        print(f"OK lớp {grade} {sid:24s} {'[rà soát]' if reviewed else '[gốc]':9s} bài={s['lessons']:3d} tiết={s['totalPeriods']:3d} tuần={s['weeks']:2d} tích hợp={s['integrations']}")
    index.sort(key=lambda x: (x["grade"], x["subjectId"]))
    # Tệp KHDH gộp cả khối (Phụ lục 2): ưu tiên bản "(đã rà soát)"
    grade_docs = {}
    for f in grade_files:
        base = os.path.basename(f)
        g = re.search(r"khối\s*(\d)", base, re.I)
        if not g:
            continue
        grade = int(g.group(1))
        reviewed = "rà soát" in base.lower()
        if grade not in grade_docs or (reviewed and not grade_docs[grade][1]):
            grade_docs[grade] = (f, reviewed)
    grade_attachments = {str(g): attach(f, g, "KHDH ca khoi") for g, (f, _) in sorted(grade_docs.items())}
    with open(os.path.join(OUT_DIR, "index.json"), "w", encoding="utf-8") as fh:
        json.dump({"schoolYear": SCHOOL_YEAR, "generatedAt": IMPORT_DATE, "items": index,
                   "gradeAttachments": grade_attachments}, fh, ensure_ascii=False, indent=1)
    print(f"Đã chép {len(index) + len(grade_attachments)} tệp Word vào {DOCS_DIR}")
    print(f"Đã ghi {len(index)} tệp JSON vào {OUT_DIR}")


if __name__ == "__main__":
    main()
